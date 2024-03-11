import os

from docx import Document

class Merger:

    def merge(self, files, input_path, output_path, input_type, output_type):

        self.input_path = input_path
        self.output_path = output_path
        self.input_type = input_type
        self.output_type = output_type
        self.files = files
        
        if input_type == '.txt' and output_type == '.docx':
            self.merge_txt_to_docx()

    def merge_txt_to_docx(self):
        # Initialize a Document object to store merged content
        doc = Document()

        # Iterate over each text file
        for txt_file in self.files:
            txt_path = os.path.join(self.input_path, txt_file)

            # Read the content of the text file
            with open(txt_path, 'r', encoding='utf-8') as f:
                text = f.read()

            # Add the content of the text file to the Document
            doc.add_paragraph(f"[{os.path.basename(txt_path)}]")
            doc.add_paragraph(text)
            doc.add_paragraph('\n')

        # Save the merged content to a DOCX file
        doc.save(os.path.join(self.output_path, 'MERGED_FILE.docx'))


    def merge_txt_to_txt(self):
        #TODO
        pass

    def merge_docx_to_txt(self):
        #TODO
        pass

    def merge_docx_to_docx(self):
        #TODO
        pass
