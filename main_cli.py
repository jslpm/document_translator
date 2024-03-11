import os
import PyPDF2
import sys
from colorama import Fore, Style
from docx import Document
from googletrans import Translator

# Example usage
folder_path = "/home/jl-dimec/Downloads/referencias"
out_converted_folder = "txt_files"
out_translated_folder = "trans_files"

def convert_pdf_to_txt(pdf_file, txt_file):
    # Open the PDF file
    with open(pdf_file, 'rb') as pdf_file:
        # Create a PDF reader object
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        # Initialize an empty string to store extracted text
        text = ''
    
        # Iterate through each page of the PDF
        for page_num in range(len(pdf_reader.pages)):
            # Extract text from the page
            page_text = pdf_reader.pages[page_num].extract_text()
            
            # Append the extracted text to the overall text
            text += page_text
        
        # Write the extracted text to a text file
        with open(txt_file, 'w', encoding='utf-8') as txt_file:
            txt_file.write(text)

def convert_all_pdfs_to_txt(folder_path, output_folder):
    # Filter out only the PDF files
    pdf_files = get_list_of_selected_files(folder_path, ".pdf")
    
    # Convert each PDF file to text
    for i, pdf_file in enumerate(pdf_files):
        print(f"{Fore.GREEN}Converting file {i+1}/{len(pdf_files)}{Style.RESET_ALL}  {pdf_file}")
        # Construct paths for input PDF and output text files
        pdf_path = os.path.join(folder_path, pdf_file)
        txt_path = os.path.join(folder_path, output_folder, os.path.splitext(pdf_file)[0] + '.txt')
        
        # Convert the PDF to text
        convert_pdf_to_txt(pdf_path, txt_path)

def get_list_of_selected_files(folder_path, file_extesion):
    # List all files in the folder
    files = os.listdir(folder_path)
    
    # Filter out only the PDF files
    selected_files = [file for file in files if file.endswith(file_extesion)]

    return selected_files

def create_folder(folder_path):
    # Create a new directory if it doesn't exist
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
        print(f"Folder '{folder_path}' created successfully.")
    else:
        print(f"{Fore.YELLOW}Folder '{folder_path}' already exists.{Style.RESET_ALL}")
        sys.exit()

def remove_line_breaks(input_file, output_file):
    # Read the contents of the input file
    with open(input_file, 'r', encoding='utf-8') as f:
        text = f.read()

    # Remove line breaks from the text
    text = text.replace('\n', ' ')

    # Write the modified text back to the output file
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(text)

def remove_line_breaks_all_files(folder_path):
    # Filter out only the TXT files
    txt_files = get_list_of_selected_files(folder_path, ".txt")

    # Convert each TXT file to text
    for i, txt_file in enumerate(txt_files):
        print(f"{Fore.GREEN}Removing break lines file {i+1}/{len(txt_files)}{Style.RESET_ALL}  {txt_file}")
        
        # Construct paths for input TXT and output text files
        txt_path = os.path.join(folder_path, txt_file)
        txt_out_path = os.path.join(folder_path, txt_file)
        
        # Convert the TXT to text
        remove_line_breaks(txt_path, txt_out_path)

def translate_all_files(folder_path, output_folder):

    txt_files = get_list_of_selected_files(folder_path, '.txt')

    for i, txt_file in enumerate(txt_files):
        print(f"{Fore.GREEN}Translating file {i+1}/{len(txt_files)}{Style.RESET_ALL}  {txt_file}")
        
        # Construct paths for input PDF and output text files
        txt_path = os.path.join(folder_path, txt_file)
        txt_out_path = os.path.join(folder_path, output_folder, txt_file)
        
        translate_text_file(txt_path, txt_out_path)


def translate_text_file(input_file, output_file, source_lang='en', target_lang='es'):
    # Read the input text from the file
    with open(input_file, 'r', encoding='utf-8') as f:
        text = f.read()

    # Create a translator object
    translator = Translator(service_urls=['translate.googleapis.com'])

    if len(text.strip()) == 0:
        print(f"{Fore.YELLOW}Empty file!{Style.RESET_ALL}  {os.path.basename(input_file)}")
    else:
        # Translate the text
        translated_text = translator.translate(text, src=source_lang, dest=target_lang)
        # Write the translated text to the output file
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(translated_text.text)

def merge_txt_files_to_docx(folder_path, output_docx):
    # Initialize a Document object to store merged content
    doc = Document()

    # Filter out only the text files
    txt_files = get_list_of_selected_files(folder_path, ".txt")

    # Iterate over each text file
    for txt_file in txt_files:
        txt_path = os.path.join(folder_path, txt_file)
        
        # Read the content of the text file
        with open(txt_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # Add the content of the text file to the Document
        doc.add_paragraph(f"[{os.path.basename(txt_path)}]")
        doc.add_paragraph(text)
        doc.add_paragraph('\n')

    # Save the merged content to a DOCX file
    doc.save(output_docx)


# Create folder for converting pdf to txt
create_folder(os.path.join(folder_path, out_converted_folder))

# Convert pdfs to txt files
convert_all_pdfs_to_txt(folder_path, out_converted_folder)

# Ask for modifing files (remove innecesary text)
print(f"Modify files for translation. When done, confirm writing {Fore.GREEN}[OK]{Style.RESET_ALL} to continue or {Fore.RED}[EXIT]{Style.RESET_ALL} to finish.")
continue_flag = input("> ")
confirmation_flags = ["OK", "EXIT"]

# Check for entered command
while continue_flag.upper() not in confirmation_flags:
    print(f"{Fore.YELLOW}Sorry, command not recognize.{Style.RESET_ALL}")
    print(f"Modify files for translation. When done, confirme writing {Fore.GREEN}[OK]{Style.RESET_ALL} to continue or {Fore.RED}[EXIT]{Style.RESET_ALL} to finish.")
    continue_flag = input("> ")

if continue_flag.upper() == confirmation_flags[0]:
    remove_line_breaks_all_files(os.path.join(folder_path, out_converted_folder))
    create_folder(os.path.join(folder_path, out_translated_folder))
    translate_all_files(os.path.join(folder_path, out_converted_folder), os.path.join(folder_path, out_translated_folder))
    merge_txt_files_to_docx(os.path.join(folder_path, out_translated_folder), os.path.join(folder_path, out_translated_folder, 'ALL_TEXT.docx'))
else:
    sys.exit()